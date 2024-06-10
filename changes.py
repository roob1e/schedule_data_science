from bs4 import BeautifulSoup
import requests
import re
import os
import doc2docx
from docx import Document
import pandas as pd
from datetime import datetime
import telebot

bot = telebot.TeleBot(Your token here)

global msg

def changes():
    link = 'http://gtec-bks.by/tmp/zamena/'
    page = requests.get('http://gtec-bks.by/stud/dnevnoe')
    soup = BeautifulSoup(page.text, "html.parser")
    print(page.status_code)

    date_ = soup.find_all('a', string=re.compile('Изменения в расписании'))
    for el in date_:
        str_ = str(el)
        break
        
    print(str_)
    changes_date = ''

    for symbol in str_:
        if symbol.isnumeric():
            changes_date += symbol
        elif changes_date.__len__() >= 8:
            break
    print(changes_date)

    #Постановка точек в дате и формировании датной ссылки
    body = changes_date[:2] + '.' + changes_date[2:4] + '.' + changes_date[4:]
    date = datetime.strptime(body, '%d.%m.%Y')
    day_of_week = date.weekday()
    link += f'{body}.doc'
    print(link) 

    response = requests.get(link) 
    #print(response)
        
    if response.status_code == 200:
        file_name = link.replace('http://gtec-bks.by/tmp/zamena/', '')
        print(file_name)
        
        if not os.path.exists('C:\CHANGES\docs'):
            os.makedirs('C:\CHANGES\docs')
            
        with open(f'C:\CHANGES\docs/{file_name}', 'wb') as file:
            file.write(response.content)
    else:
        print('Ошибка загрузки файла')
    print(file_name)

    
    file_path = 'C:\CHANGES\docs' + f"\{file_name}"
    print(file_path)

    doc2docx.main()

    docx_path = 'C:\CHANGES\doc_convert\docs.docx'
    pd.set_option('display.max_rows', 200)
    pd.set_option('display.max_columns', 10)
    # Открываем документ
    doc = Document(docx_path)

    # Получаем первую таблицу в документе
    table = doc.tables[0]

    for paragraphs in doc.paragraphs:
        if 'НАД ЧЕРТОЙ' in paragraphs.text:
            up_line = True
            break
        elif 'ПОД ЧЕРТОЙ' in paragraphs.text:
            up_line = False
            break
        
    # Создаем список для хранения данных
    data = []

    # Проходим по всем строкам таблицы
    for row in table.rows:
        # Создаем список для хранения данных в строке
        row_data = []
        # Проходим по всем ячейкам в строке
        for cell in row.cells:
            # Добавляем текст из ячейки в список данных строки
            row_data.append(cell.text)
        # Добавляем список данных строки в основной список данных
        data.append(row_data)

    # Создаем DataFrame из списка данных
    df = pd.DataFrame(data)

    # Выводим DataFrame
    df.to_csv('data.csv', index=False)
    df.columns = ['Группа', "Пара", "Предмет", "Изменения", "Аудитория"]
    print(df.columns.tolist())
    df['Группа'] = df['Группа'].str.replace(' ', '', regex=True).replace('\n', '', regex=True)
    selected_rows = df.loc[df['Группа'] == 'П-21']
    print(selected_rows)
    if selected_rows.empty == True:
        print('Изменений нет!')
        no_changes = True
    else:
        no_changes = False

    schedule_monday_up = [
        ['Понедельник', 1, '09:00 - 10:40', 'Язык программирования Java: разр. алгоритмов', 'Васьковцова Н.С.', '404'],
        ['Понедельник', 2, '10:50 - 12:30', 'Основы алгоритмизации и программирования', 'Васьковцова Н.С.', '404'], 
        ['Понедельник', 3, '13:00 - 14:40', 'Физ. культура и здоровье / Основы соц-гуманитарных наук', 'Преподаватель / Фролова Т.Н.', 'Сп.з. / 213'], 
        ['Понедельник', 4, '14:50 - 16:30', 'Технология разработки ПО', 'Курилина Е.М.', 405]     
    ]
    schedule_tuesday_up = [
        ['Вторник', 1, '09:00 - 10:40', 'Математика', 'Фоминых Е.И.', '312'],
        ['Вторник', 2, '10:50 - 12:30', 'Иностранный язык (проф. лексика)', 'Орлова Я.Е. / Соловьева Л.И.', '216'], 
        ['Вторник', 3, '13:00 - 14:40', 'Технология разработки ПО', 'Курилина Е.М.', '405'], 
        ['Вторник', 4, '14:50 - 15:35', 'Час куратора', 'Фролова Т.Н.', '213']    
    ]
    schedule_wednesday_up = [
        ['Среда', 1, '09:00 - 10:40', 'Математика', 'Фоминых Е.И.', '312'],
        ['Среда', 2, '10:50 - 12:30', 'Инструментальное ПО', 'Алова Е.В.', '305'], 
        ['Среда', 3, '13:00 - 14:40', 'Физ. культура и здоровье / Основы соц-гуманитарных наук', 'Преподаватель / Фролова Т.Н.', 'Сп.з. / 213'],
        ['Среда', 4, '14:50 - 16:30', 'Стандартиз. и сертифик. ПО', 'Железная С.В.', '419'],  
    ]
    schedule_thursday_up = [
        ['Четверг', 2, '10:50 - 12:30', 'Теория вероятности и математич. статистика', 'Фоминых Е.И.', '312'], 
        ['Четверг', 3, '13:00 - 14:40', 'Язык программирования Java: разр. алгоритмов', 'Васьковцова Н.С.', '404'], 
        ['Четверг', 4, '14:50 - 16:30', 'Арифм.-лог. осн. вычисл. техники', 'Рудько И.Д.', '309']  
    ]
    schedule_friday_up = [
        ['Пятница', 1, '09:00 - 10:40', 'Арифм.-лог. осн. вычисл. техники', 'Рудько И.Д.', '309'],
        ['Пятница', 2, '10:50 - 12:30', 'Технология разработки ПО', 'Курилина Е.М.', '405'], 
        ['Пятница', 3, '13:00 - 14:40', 'Физ. культура и здоровье / Основы соц-гуманитарных наук', 'Преподаватель / Фролова Т.Н.', 'Сп.з. / 213']   
    ]
    schedule_saturday_up = [
        ['Суббота', 1, '08:00 - 09:40', 'Язык программирования Java: разр. алгоритмов', 'Васьковцова Н.С.', '404'],
        ['Суббота', 2, '09:50 - 11:30', 'Основы алгоритмизации и программирования', 'Васьковцова Н.С.', '404'], 
        ['Суббота', 3, '12:00 - 13:40', 'Математика', 'Фоминых Е.И.', '312']   
    ]







    schedule_monday_down = [
        ['Понедельник', 1, '09:00 - 10:40', 'Язык программирования Java: разр. алгоритмов', 'Васьковцова Н.С.', '404'],
        ['Понедельник', 2, '10:50 - 12:30', 'Основы алгоритмизации и программирования', 'Васьковцова Н.С.', '404'], 
        ['Понедельник', 3, '13:00 - 14:40', 'Физ. культура и здоровье / Основы соц-гуманитарных наук', 'Преподаватель / Фролова Т.Н.', 'Сп.з. / 213'], 
        ['Понедельник', 4, '14:50 - 16:30', 'Технология разработки ПО', 'Курилина Е.М.', 405]     
    ]
    schedule_tuesday_down = [
        ['Вторник', 1, '09:00 - 10:40', 'Математика', 'Фоминых Е.И.', '312'],
        ['Вторник', 2, '10:50 - 12:30', 'Иностранный язык (проф. лексика)', 'Орлова Я.Е. / Соловьева Л.И.', '216'], 
        ['Вторник', 3, '13:00 - 14:40', 'Технология разработки ПО', 'Курилина Е.М.', '405'], 
        ['Вторник', 4, '14:50 - 15:35', 'Час куратора', 'Фролова Т.Н.', '213']    
    ]
    schedule_wednesday_down = [
        ['Среда', 1, '09:00 - 10:40', 'Стандартиз. и сертифик. ПО', 'Железная С.В.', '419'],
        ['Среда', 2, '10:50 - 12:30', 'Инструментальное ПО', 'Алова Е.В.', '305'], 
        ['Среда', 3, '13:00 - 14:40', 'Физ. культура и здоровье / Основы соц-гуманитарных наук', 'Преподаватель / Фролова Т.Н.', 'Сп.з. / 213'] 
    ]
    schedule_thursday_down = [
        ['Четверг', 2, '10:50 - 12:30', 'Теория вероятности и математич. статистика', 'Фоминых Е.И.', '312'], 
        ['Четверг', 3, '13:00 - 14:40', 'Основы алгоритмизации и программирования', 'Васьковцова Н.С.', '404'], 
        ['Четверг', 4, '14:50 - 16:30', 'Арифм.-лог. осн. вычисл. техники', 'Рудько И.Д.', '309']  
    ]
    schedule_friday_down = [
        ['Пятница', 1, '09:00 - 10:40', 'Защ. насел. и террит. от ЧС', 'Лепшая Я.В.', '414'],
        ['Пятница', 2, '10:50 - 12:30', 'Математика', 'Фоминых Е.И.', '312'], 
        ['Пятница', 3, '13:00 - 14:40', 'Физ. культура и здоровье / Основы соц-гуманитарных наук', 'Преподаватель / Фролова Т.Н.', 'Сп.з. / 213']   
    ]
    schedule_saturday_down = [
        ['Суббота', 1, '08:00 - 09:40', 'Язык программирования Java: разр. алгоритмов', 'Васьковцова Н.С.', '404'],
        ['Суббота', 2, '09:50 - 11:30', 'Основы алгоритмизации и программирования', 'Васьковцова Н.С.', '404'], 
        ['Суббота', 3, '12:00 - 13:40', 'Математика', 'Фоминых Е.И.', '312']   
    ]









    days = ["Понедельник", "Вторник", "Среда", "Четверг", "Пятница", "Суббота", "Воскресенье"]
    if up_line:
        match day_of_week:
            case 0:
                df_ = schedule_monday_up
            case 1:
                df_ = schedule_tuesday_up
            case 2:
                df_ = schedule_wednesday_up
            case 3:
                df_ = schedule_thursday_up
            case 4: 
                df_ = schedule_friday_up
            case 5:
                df_ = schedule_saturday_up
    else:
        match day_of_week:
            case 0:
                df_ = schedule_monday_down
            case 1:
                df_ = schedule_tuesday_down
            case 2:
                df_ = schedule_wednesday_down
            case 3:
                df_ = schedule_thursday_down
            case 4: 
                df_ = schedule_friday_down
            case 5:
                df_ = schedule_saturday_down    
    actual_schedule_df = pd.DataFrame(df_, columns=['День', 'Пара', 'Время', 'Предмет', 'Преподаватель', 'Аудитория'])
    print(actual_schedule_df)
    global msg
    msg = ''
    if no_changes == True:
        for i in range(len(actual_schedule_df.index)):
            msg = msg + f'{actual_schedule_df['Пара'][i]} - {actual_schedule_df['Предмет'][i]} - {actual_schedule_df['Преподаватель'][i]} - {actual_schedule_df['Аудитория'][i]}\n\n' 
    elif no_changes == False:
        range_index = len(selected_rows.index)
        index_r = max(range_index, len(actual_schedule_df.index))
            
        for i in range(index_r):
            if i in selected_rows.index and int(actual_schedule_df['Пара'][i]) == int(selected_rows['Пара'][i]):
                msg = msg + f'{selected_rows["Пара"][i]} - {selected_rows['Изменения'][i]}\n'
            else:
                msg = msg + f'{actual_schedule_df['Пара'][i]} - {actual_schedule_df['Предмет'][i]} - {actual_schedule_df['Преподаватель'][i]} - {actual_schedule_df['Аудитория'][i]}\n\n'
    return f'Расписание на {body} с учётом изменений:\n{msg}'


@bot.message_handler(commands=['pairs'])
def send_schedule(message):
    chat_id = message.chat.id
    bot.send_message(chat_id, changes())
    
if __name__=='__main__':
    print('Бот запущен')
    bot.polling()   
